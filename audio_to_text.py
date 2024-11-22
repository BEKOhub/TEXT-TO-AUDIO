# %%
# Aactivate audiodoc\Scripts\activate


#%%

# Configurer FFmpeg dans Python

from pydub import AudioSegment

# Spécifiez le chemin vers FFmpeg et FFprobe
AudioSegment.converter = "C:/Paths/ffmpeg/ffmpeg-7.1-essentials_build/bin/ffmpeg.exe"
AudioSegment.ffprobe = "C:/Paths/ffmpeg/ffmpeg-7.1-essentials_build/bin/ffprobe.exe"

# Convertir un fichier .m4a en .wav
audio = AudioSegment.from_file("./audiofile/Nanterre - La Folie 2.m4a", format="m4a")
audio.export("./audiofile/Nanterre - La Folie 2.wav", format="wav")
print("Conversion réussie.")

# %% to 1 min Segments

from pydub import AudioSegment

# Charger le fichier audio
audio = AudioSegment.from_wav("./audiofile/Nanterre - La Folie 2.wav")

# Découper en segments de 60 secondes
segment_length = 60 * 1000  # 60 secondes en millisecondes
segments = [audio[i:i + segment_length] for i in range(0, len(audio), segment_length)]

# Exporter les segments
for idx, segment in enumerate(segments):
    segment.export(f"./audiofile/segments/segment_{idx}.wav", format="wav")
    print(f"Segment {idx} exporté.")


# %%

import os
from concurrent.futures import ThreadPoolExecutor
from docx import Document
import speech_recognition as sr

# Configuration
segments_folder = "./audiofile/segments/"
output_file = "./transcriptions.docx"

# Vérifier si le dossier existe
if not os.path.exists(segments_folder):
    print(f"Le dossier '{segments_folder}' n'existe pas. Vérifiez le chemin.")
    exit()

# Initialisation
doc = Document()
doc.add_heading("Transcriptions Audio", level=1)
recognizer = sr.Recognizer()

# Fonction pour transcrire un fichier
def transcribe_segment(segment_file):
    segment_path = os.path.join(segments_folder, segment_file)
    try:
        with sr.AudioFile(segment_path) as source:
            audio_data = recognizer.record(source)
            transcription = recognizer.recognize_google(audio_data, language="fr-FR")
            return f"Fichier : {segment_file}\nTranscription : {transcription}\n\n"
    except Exception as e:
        return f"Fichier : {segment_file}\nErreur : {e}\n\n"

# Traiter les fichiers en parallèle
with ThreadPoolExecutor() as executor:
    segment_files = [f for f in sorted(os.listdir(segments_folder)) if f.endswith(".wav")]
    results = list(executor.map(transcribe_segment, segment_files))

# Ajouter les résultats au document Word
for result in results:
    doc.add_paragraph(result)

# Sauvegarder le document Word
doc.save(output_file)
print(f"Transcriptions sauvegardées dans : {output_file}")




# %%
